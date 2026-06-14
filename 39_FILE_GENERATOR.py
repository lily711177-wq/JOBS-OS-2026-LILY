#!/usr/bin/env python3
"""
Professional Resume & Cover Letter Generator v3 — JOBS OS 2026
Fixed: re-flows hard-wrapped source text into natural paragraphs.
"""

import os, re, sys

ROOT = "/mnt/e/JOBS OS 2026/APPLICATION_PACKAGES"
OUTDIR = "OUTPUT"

REAL_NAME = "[NAME]"
REAL_PHONE = "[PHONE]"
REAL_EMAIL = "[EMAIL]"
REAL_LINKEDIN = "[LINKEDIN]"

SECTION_HEADERS = {"PROFESSIONAL SUMMARY", "CORE COMPETENCIES",
    "PROFESSIONAL EXPERIENCE", "KEY RESULTS", "EDUCATION",
    "TECHNICAL PROFICIENCY", "CORE COMPETENCIAS"}

def detect_format(md_content):
    ats_line = ""
    for line in md_content.split('\n')[:10]:
        if 'ATS:' in line:
            ats_line = line.lower()
            break
    # All companies get DOCX — visual quality is proven, every ATS accepts it
    return "docx"

def replace_placeholders(text):
    text = text.replace("ARYAN S.", REAL_NAME)
    text = text.replace("[phone]", REAL_PHONE)
    text = text.replace("[email]", REAL_EMAIL)
    text = text.replace("[linkedin]", REAL_LINKEDIN)
    text = text.replace("linkedin.com/in/aryans", REAL_LINKEDIN)
    text = text.replace("www.linkedin.com/in/aryans", REAL_LINKEDIN)
    return text

def extract_sections(md_content):
    m = re.search(r'## RESUME.*?```\n(.*?)```', md_content, re.DOTALL)
    if not m:
        m = re.search(r'## \d+\.\s*RESUME.*?```\n(.*?)```', md_content, re.DOTALL)
    resume_text = m.group(1).strip() if m else ""
    cover_text = ""
    m = re.search(r'## COVER LETTER\n+(.*?)(?=\n##|\Z)', md_content, re.DOTALL)
    if m: cover_text = m.group(1).strip()
    m2 = re.search(r'## \d+\.\s*COVER LETTER.*?\n(.*?)(?=\n##|\Z)', md_content, re.DOTALL)
    if m2 and not cover_text: cover_text = m2.group(1).strip()
    if cover_text:
        cover_text = re.sub(r'\*\*Target:.*?\*\*.*\n?', '', cover_text)
        cover_text = re.sub(r'\*\*Max.*?\*\*.*\n?', '', cover_text)
        cover_text = re.sub(r'^---\s*\n?', '', cover_text, flags=re.MULTILINE)
        cover_text = cover_text.strip()
    return resume_text, cover_text


def classify_line(ls, idx, total_lines, lines):
    if not ls: return "empty"
    is_bullet = ls.startswith('\u2022') or ls.startswith('-')
    if is_bullet: return "bullet"
    if idx == 0: return "name"
    if any(x in ls.lower() for x in ['@', 'phone', 'linkedin', 'vancouver, bc']): return "contact"
    if ls.strip() in SECTION_HEADERS: return "section"
    if ls == ls.upper() and len(ls) > 3 and len(ls) < 60: return "subheader"
    if idx > 0:
        prev = lines[idx-1].strip() if lines[idx-1] else ""
        if prev in SECTION_HEADERS: return "lead_para"
        if prev == prev.upper() and len(prev) > 3 and len(prev) < 60 and prev not in SECTION_HEADERS:
            return "lead_para"
    title_words = ['chief', 'director', 'manager', 'head of', 'lead ', 'officer', 'general manager']
    if any(w in ls.lower() for w in title_words) and len(ls) < 80: return "title"
    if ls.startswith('['): return "company"
    return "body"


def reflow_lines(lines):
    """Reflow hard-wrapped text into flowing paragraphs.
    Bullet continuations (body/title/lead_para after a bullet) get merged into the bullet.
    Consecutive flowable lines get merged into single paragraphs."""
    result = []
    buf = []
    buf_class = None
    for idx, line in enumerate(lines):
        ls = line.strip()
        cls = classify_line(ls, idx, len(lines), lines)

        # Body/title/lead_para after a bullet = bullet continuation
        if buf_class == "bullet" and cls in ("body", "title", "lead_para"):
            buf.append(ls)
            continue

        # Flowable lines merge together
        if cls in ("body", "lead_para") and buf_class in ("body", "lead_para"):
            buf.append(ls)
            continue

        if buf:
            result.append((buf_class, ' '.join(buf)))
            buf = []
            buf_class = None

        buf.append(ls)
        buf_class = cls

    if buf:
        result.append((buf_class, ' '.join(buf)))
    return result


# ═══════════════════════════════════════════════════════════════
#  DOCX GENERATOR
# ═══════════════════════════════════════════════════════════════

def gen_docx(pkg_dir, slug, resume_text, cover_text):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    od = os.path.join(pkg_dir, OUTDIR)
    os.makedirs(od, exist_ok=True)

    def border(p):
        pPr = p._p.get_or_add_pPr()
        b = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '555555')
        b.append(bottom); pPr.append(b)

    def mp(text, size=10, bold=False, align=None, before=0, after=2):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(size); r.font.name = 'Calibri'; r.bold = bold
        pf = p.paragraph_format; pf.space_before = Pt(before); pf.space_after = Pt(after)
        if align: pf.alignment = align
        return p

    def mb(text):
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        r = p.add_run(text)
        r.font.size = Pt(10); r.font.name = 'Calibri'
        pf = p.paragraph_format; pf.space_before = Pt(0.5); pf.space_after = Pt(0.5)
        pf.left_indent = Inches(0.3)
        return p

    doc = Document()
    style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2); style.paragraph_format.space_before = Pt(0)
    for s in doc.sections:
        s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.8); s.right_margin = Inches(0.8)

    raw = [l.rstrip() for l in resume_text.split('\n')]
    grouped = reflow_lines(raw)

    for cls, text in grouped:
        if cls == "empty": mp('', size=6)
        elif cls == "name": mp(text, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
        elif cls == "contact": mp(text, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
        elif cls == "section": mp(text, size=11, bold=True, before=8, after=3); border(doc.paragraphs[-1])
        elif cls == "subheader": mp(text, size=10, bold=True, before=6, after=1)
        elif cls == "lead_para": mp(text, size=10, before=2, after=3)
        elif cls == "title": mp(text, size=10, bold=True, before=4, after=1)
        elif cls == "company": mp(text, size=9, after=2)
        elif cls == "bullet": mb(text.lstrip('\u2022- ').strip())
        else: mp(text, size=10, after=1)

    rp = os.path.join(od, f"{slug}_RESUME.docx")
    doc.save(rp)

    doc = Document()
    s = doc.styles['Normal']; s.font.name = 'Calibri'; s.font.size = Pt(10.5)
    for sec in doc.sections:
        sec.top_margin = Inches(1.0); sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0); sec.right_margin = Inches(1.0)

    in_body = False; sig_done = False
    for line in cover_text.split('\n'):
        ls = line.strip()
        if not ls: mp('', size=8); continue
        if ls.startswith('**') and ':' in ls: continue
        if 'Dear' in ls: mp(ls, size=10.5, before=4, after=6); in_body = True; continue
        if ls.startswith('Best') or ls.startswith('Sincerely') or ls.startswith('Warmly'):
            mp(ls, size=10.5, before=10); sig_done = True; continue
        if sig_done: mp(ls, size=10.5); continue
        if in_body: mp(ls, size=10.5, after=4)
        else: mp(ls, size=10.5)

    cp = os.path.join(od, f"{slug}_COVER_LETTER.docx")
    doc.save(cp)
    return rp, cp


# ═══════════════════════════════════════════════════════════════
#  PDF GENERATOR — HTML→WeasyPrint with reflow
# ═══════════════════════════════════════════════════════════════

def gen_pdf(pkg_dir, slug, resume_text, cover_text):
    from weasyprint import HTML

    od = os.path.join(pkg_dir, OUTDIR)
    os.makedirs(od, exist_ok=True)

    def text_to_html(text_lines):
        raw = text_lines.split('\n')
        lines = [l.rstrip() for l in raw]
        grouped = reflow_lines(lines)

        css = """
        @page { size: Letter; margin: 0.7in 0.8in 0.7in 0.8in; }
        body { font-family: 'DejaVu Sans', sans-serif; font-size: 10pt; color: #222; margin: 0; padding: 0; }
        .name { font-size: 16pt; font-weight: 700; text-align: center; margin: 0 0 2pt 0; }
        .contact { font-size: 9pt; text-align: center; margin: 0 0 8pt 0; color: #444; }
        .section { font-size: 11pt; font-weight: 700; margin: 12pt 0 5pt 0; padding: 0 0 3pt 0; border-bottom: 0.5pt solid #999; }
        .subheader { font-size: 10pt; font-weight: 700; margin: 7pt 0 2pt 0; }
        .lead_para { margin: 0 0 6pt 0; line-height: 1.45; }
        .title { font-size: 10pt; font-weight: 700; margin: 6pt 0 2pt 0; }
        .company { font-size: 9pt; margin: 0 0 3pt 0; color: #555; }
        .body { margin: 0 0 4pt 0; line-height: 1.4; }
        .bullet { margin: 1.5pt 0 1.5pt 18pt; padding-left: 6pt; text-indent: -6pt; line-height: 1.4; }
        .spacer { height: 3pt; }
        """

        html_parts = ['<html><head><meta charset="utf-8"><style>' + css + '</style></head><body>']

        for cls, text in grouped:
            if cls == "empty":
                html_parts.append('<div class="spacer"></div>')
            elif cls == "name":
                html_parts.append(f'<div class="name">{text}</div>')
            elif cls == "contact":
                html_parts.append(f'<div class="contact">{text}</div>')
            elif cls == "section":
                html_parts.append(f'<div class="section">{text}</div>')
            elif cls == "subheader":
                html_parts.append(f'<div class="subheader">{text}</div>')
            elif cls == "lead_para":
                html_parts.append(f'<div class="lead_para">{text}</div>')
            elif cls == "title":
                html_parts.append(f'<div class="title">{text}</div>')
            elif cls == "company":
                html_parts.append(f'<div class="company">{text}</div>')
            elif cls == "bullet":
                t = text.lstrip('\u2022- ').strip()
                html_parts.append(f'<div class="bullet">{t}</div>')
            elif cls == "body":
                html_parts.append(f'<div class="body">{text}</div>')
            else:
                html_parts.append(f'<div style="margin:1pt 0">{text}</div>')

        html_parts.append('</body></html>')
        return '\n'.join(html_parts)

    def cover_to_html(text):
        css = """
        @page { size: Letter; margin: 1in; }
        body { font-family: 'DejaVu Sans', sans-serif; font-size: 10.5pt; color: #222; margin: 0; padding: 0; line-height: 1.55; }
        .greeting { margin: 8pt 0 12pt 0; }
        .para { margin: 0 0 10pt 0; }
        .closing { margin: 18pt 0 0 0; }
        .sig { margin: 0; }
        .spacer { height: 6pt; }
        """
        html_parts = ['<html><head><meta charset="utf-8"><style>' + css + '</style></head><body>']
        in_body = False; sig_done = False
        for line in text.split('\n'):
            ls = line.strip()
            if not ls:
                html_parts.append('<div class="spacer"></div>')
                continue
            if ls.startswith('**') and ':' in ls:
                continue
            if 'Dear' in ls:
                html_parts.append(f'<div class="greeting">{ls}</div>')
                in_body = True; continue
            if ls.startswith('Best') or ls.startswith('Sincerely') or ls.startswith('Warmly'):
                html_parts.append(f'<div class="closing">{ls}</div>')
                sig_done = True; continue
            if sig_done:
                html_parts.append(f'<div class="sig">{ls}</div>')
            elif in_body:
                html_parts.append(f'<div class="para">{ls}</div>')
            else:
                html_parts.append(f'<div class="para">{ls}</div>')
        html_parts.append('</body></html>')
        return '\n'.join(html_parts)

    rp = os.path.join(od, f"{slug}_RESUME.pdf")
    HTML(string=text_to_html(resume_text)).write_pdf(rp)
    cp = os.path.join(od, f"{slug}_COVER_LETTER.pdf")
    HTML(string=cover_to_html(cover_text)).write_pdf(cp)
    return rp, cp


# ═══════════════════════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════════════════════

def main():
    total = 0; errors = []
    for entry in sorted(os.listdir(ROOT)):
        pkg_dir = os.path.join(ROOT, entry)
        if not os.path.isdir(pkg_dir): continue
        md_path = os.path.join(pkg_dir, "APPLICATION_PACKAGE.md")
        if not os.path.exists(md_path): continue

        with open(md_path) as f: md = f.read()
        resume_text, cover_text = extract_sections(md)
        if not resume_text:
            errors.append(f"{entry}: no resume text"); continue

        resume_text = replace_placeholders(resume_text)
        cover_text = replace_placeholders(cover_text)

        fmt = detect_format(md)

        out_dir = os.path.join(pkg_dir, OUTDIR)
        if os.path.exists(out_dir):
            for fn in os.listdir(out_dir):
                try: os.remove(os.path.join(out_dir, fn))
                except: pass

        try:
            if fmt == "pdf":
                gen_pdf(pkg_dir, entry, resume_text, cover_text)
                print(f"  {chr(10003)} {entry}  [PDF]")
            else:
                gen_docx(pkg_dir, entry, resume_text, cover_text)
                print(f"  {chr(10003)} {entry}  [DOCX]")
            total += 1
        except Exception as e:
            errors.append(f"{entry}: {e}")
            print(f"  {chr(10007)} {entry}: {e}")

    print(f"\n{'='*60}")
    print(f"  Generated: {total} packages")
    if errors:
        print(f"  Errors: {len(errors)}")
        for e in errors: print(f"    - {e}")
    print(f"{'='*60}")

if __name__ == "__main__":
    main()
